from io import BytesIO

import pandas as pd
from docx import Document
from docx.shared import Inches

from src.eda import categorical_summary, column_summary, numeric_summary, outlier_summary


def _add_table(document: Document, dataframe: pd.DataFrame, max_rows: int = 30) -> None:
    frame = dataframe.head(max_rows).fillna("").astype(str)
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    for index, column in enumerate(frame.columns):
        table.rows[0].cells[index].text = str(column)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def generate_report(df: pd.DataFrame, file_name: str, field_types: dict[str, list[str]]) -> bytes:
    document = Document()
    document.add_heading("DataInsight Agent 分析报告", 0)
    document.add_paragraph(f"文件：{file_name}")
    document.add_paragraph(f"数据规模：{len(df)} 行 × {len(df.columns)} 列")

    document.add_heading("数据概览与缺失值分析", level=1)
    columns = column_summary(df)
    _add_table(document, columns)
    high_missing = columns.loc[columns["缺失值比例"] >= 20, "字段名"].tolist()
    document.add_paragraph(f"洞察：缺失率超过 20% 的字段：{', '.join(map(str, high_missing)) or '无'}。")

    document.add_heading("重复值分析", level=1)
    document.add_paragraph(f"重复行数量：{int(df.duplicated().sum())}")

    numeric = numeric_summary(df, field_types["numeric"])
    document.add_heading("数值字段统计", level=1)
    if numeric.empty:
        document.add_paragraph("无数值字段。")
    else:
        _add_table(document, numeric)
        skewed = numeric.loc[numeric["skewness"].abs() > 1, "字段名"].tolist()
        document.add_paragraph("偏度绝对值大于 1 通常表示明显偏态；峰度越高，极端值越集中。")
        document.add_paragraph(f"明显偏态字段：{', '.join(map(str, skewed)) or '无'}。")

    outliers = outlier_summary(df, field_types["numeric"])
    document.add_heading("异常值分析（IQR）", level=1)
    if not outliers.empty:
        _add_table(document, outliers)
        top = outliers.sort_values("异常值数量", ascending=False).head(3)["字段名"].tolist()
        document.add_paragraph(f"异常值较多的字段：{', '.join(map(str, top)) or '无'}。")

    categories = categorical_summary(df, field_types["categorical"])
    document.add_heading("类别字段洞察", level=1)
    if not categories.empty:
        _add_table(document, categories)

    document.add_heading("图表说明", level=1)
    document.add_paragraph("建议结合应用中的分布图、箱型图、趋势图和相关性热力图进一步判断数据规律。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
